import streamlit as st
import pandas as pd
import io
import re
import ssl
import os
from Bio import Entrez, Medline
from docx import Document
from difflib import get_close_matches
import xml.etree.ElementTree as ET

if (not os.environ.get('PYTHONHTTPSVERIFY', '') and
        getattr(ssl, '_create_unverified_context', None)):
    ssl._create_default_https_context = ssl._create_unverified_context

st.set_page_config(page_title="PubMed Web Scanner by RTOmega", page_icon="🧬", layout="wide")

STUDY_TYPES = [
    "Adaptive Clinical Trial", "Address", "Autobiography", "Bibliography", "Biography", 
    "Books and Documents", "Case Reports", "Classical Article", "Clinical Conference", 
    "Clinical Study", "Clinical Trial", "Clinical Trial Protocol", "Clinical Trial, Phase I", 
    "Clinical Trial, Phase II", "Clinical Trial, Phase III", "Clinical Trial, Phase IV", 
    "Clinical Trial, Veterinary", "Collected Work", "Comment", "Comparative Study", 
    "Congress", "Consensus Development Conference", "Consensus Development Conference, NIH", 
    "Controlled Clinical Trial", "Corrected and Republished Article", "Dataset", "Dictionary", 
    "Directory", "Duplicate Publication", "Editorial", "Electronic Supplementary Materials", 
    "English Abstract", "Equivalence Trial", "Evaluation Study", "Expression of Concern", 
    "Festschrift", "Government Publication", "Guideline", "Historical Article", 
    "Interactive Tutorial", "Interview", "Introductory Journal Article", "Lecture", 
    "Legal Case", "Legislation", "Letter", "Meta-Analysis", "Multicenter Study", 
    "Network Meta-Analysis", "News", "Newspaper Article", "Observational Study", 
    "Observational Study, Veterinary", "Overall", "Patient Education Handout", 
    "Periodical Index", "Personal Narrative", "Portrait", "Practice Guideline", 
    "Pragmatic Clinical Trial", "Preprint", "Published Erratum", "Randomized Controlled Trial", 
    "Randomized Controlled Trial, Veterinary", "Research Support, American Recovery and Reinvestment Act", 
    "Research Support, N.I.H., Extramural", "Research Support, N.I.H., Intramural", 
    "Research Support, Non-U.S. Gov't", "Research Support, U.S. Gov't, Non-P.H.S.", 
    "Research Support, U.S. Gov't, P.H.S.", "Research Support, U.S. Gov't", 
    "Retracted Publication", "Retraction of Publication", "Review", "Scientific Integrity Review", 
    "Scoping Review", "Systematic Review", "Technical Report", "Twin Study", 
    "Validation Study", "Video-Audio Media", "Webcast"
]

def normalize_journal_name(name):
    if not isinstance(name, str):
        return ""
    name = name.lower()
    name = re.sub(r'\bthe\b', '', name)
    name = re.sub(r'[^a-z0-9 ]+', ' ', name)
    return name.strip()

def get_pmc_corresp_info(pmc_ids):
    if not pmc_ids:
        return {}
        
    Entrez.email = "pubmed_tool_web@example.com"
    clean_ids = [pid.replace("PMC", "") for pid in pmc_ids if pid]
    if not clean_ids:
        return {}
        
    try:
        handle = Entrez.efetch(db="pmc", id=",".join(clean_ids), rettype="xml")
        xml_data = handle.read()
        handle.close()
    except Exception as e:
        print("Error fetching pmc:", e)
        return {}

    try:
        root = ET.fromstring(xml_data)
    except Exception as e:
        print("Error parsing pmc xml:", e)
        return {}

    results = {}
    for article in root.findall(".//article"):
        pmcid_node = article.find(".//article-id[@pub-id-type='pmcid']")
        if pmcid_node is None:
            continue
        pmcid = pmcid_node.text
        if pmcid and not pmcid.startswith("PMC"):
            pmcid = "PMC" + pmcid
            
        corresp_email = ""
        corresp_name = ""
        
        corresp_dict = {}
        for corresp in article.findall(".//corresp"):
            cid = corresp.get("id")
            email_node = corresp.find(".//email")
            email = email_node.text if email_node is not None else ""
            if email:
                if cid:
                    corresp_dict[cid] = email
                if not corresp_email:
                    corresp_email = email
                    
        for contrib in article.findall(".//contrib[@contrib-type='author']"):
            is_corresp = False
            if contrib.get("corresp") == "yes":
                is_corresp = True
                
            email_for_this_author = ""
            for xref in contrib.findall("xref[@ref-type='corresp']"):
                rid = xref.get("rid")
                if rid in corresp_dict:
                    is_corresp = True
                    email_for_this_author = corresp_dict[rid]
                    
            if is_corresp:
                name = contrib.find("name")
                if name is not None:
                    surname = name.find("surname")
                    given = name.find("given-names")
                    s_str = surname.text if surname is not None else ""
                    g_str = given.text if given is not None else ""
                    corresp_name = f"{g_str} {s_str}".strip()
                    if email_for_this_author:
                        corresp_email = email_for_this_author
                break 
                
        if pmcid:
            results[pmcid] = {
                "name": corresp_name,
                "email": corresp_email
            }
        
    return results

def search_pubmed(query, max_results):
    Entrez.email = "pubmed_tool_web@example.com"
    try:
        handle = Entrez.esearch(db="pubmed", term=query, retmax=max_results)
        record = Entrez.read(handle)
        handle.close()
        
        ids = record["IdList"]
        if not ids:
            return pd.DataFrame()

        handle = Entrez.efetch(db="pubmed", id=",".join(ids), rettype="medline", retmode="text")
        records = Medline.parse(handle)
        
        articles = []
        for r in records:
            # Handle DOI
            doi_raw = r.get("LID", r.get("AID", ""))
            doi_link = ""
            if doi_raw and "[doi]" in doi_raw:
                clean_doi = doi_raw.split(' ')[0]
                doi_link = f"https://doi.org/{clean_doi}"

            pmid = r.get("PMID", "")
            pmid_link = f"https://pubmed.ncbi.nlm.nih.gov/{pmid}/" if pmid else ""
            
            # Handle PMCID
            pmc_raw = r.get("PMC", "")
            pmc_link = f"https://www.ncbi.nlm.nih.gov/pmc/articles/{pmc_raw}/" if pmc_raw else ""

            articles.append({
                "Select": False,
                "PMID": pmid_link, # Store URL, display ID later
                "PMCID": pmc_link,
                "Title": r.get("TI", ""),
                "First Author": r.get("AU", ["N/A"])[0],
                "Journal": r.get("JT", ""),
                "Year": r.get("DP", "N/A")[:4],
                "DOI": doi_link,
                "Article Type": "; ".join(r.get("PT", [])),
                "Keywords": "; ".join(r.get("OT", []))
            })
            
        # Fetch corresponding author info for found PMCIDs
        pmc_ids = [r.get("PMC", "") for r in records if r.get("PMC", "")]
        corresp_info = get_pmc_corresp_info(pmc_ids[:100]) # Batched up to 100 max safe check, Entrez usually handles large
        
        for a in articles:
            pmc_val = a["PMCID"]
            pmcid_clean = pmc_val.strip("/").split("/")[-1] if pmc_val else ""
            info = corresp_info.get(pmcid_clean, {})
            a["Corresp. Author Name"] = info.get("name", "N/A") if info.get("name") else "N/A"
            a["Corresp. Author Email"] = info.get("email", "N/A") if info.get("email") else "N/A"
            
        return pd.DataFrame(articles)
    except Exception as e:
        st.error(f"Error connecting to PubMed: {e}")
        return pd.DataFrame()

def process_quartiles(df, file_source):
    if file_source is None:
        df["Quartile"] = "Unknown (No File)"
        return df

    try:
        sjr = pd.read_csv(file_source, sep=';', quotechar='"', on_bad_lines='warn')
        
        title_col = next((c for c in sjr.columns if c.lower() in ["title", "journal title", "source title"]), None)
        quartile_col = next((c for c in sjr.columns if "quartile" in c.lower()), None)

        if not title_col or not quartile_col:
            df["Quartile"] = "Unknown (Column Error)"
            return df

        sjr["norm_title"] = sjr[title_col].apply(normalize_journal_name)
        quartile_map = dict(zip(sjr["norm_title"], sjr[quartile_col]))
        
        journal_names_norm = df["Journal"].apply(normalize_journal_name)
        quartiles = []
        
        for norm_name in journal_names_norm:
            if norm_name in quartile_map:
                quartiles.append(quartile_map[norm_name])
            else:
                close = get_close_matches(norm_name, quartile_map.keys(), n=1, cutoff=0.8)
                quartiles.append(quartile_map[close[0]] if close else "Unknown")

        df["Quartile"] = quartiles
        return df
    except Exception as e:
        st.warning(f"Error processing Quartile file: {e}")
        df["Quartile"] = "Error"
        return df

def to_excel(df):
    output = io.BytesIO()
    export_df = df.drop(columns=['Select'], errors='ignore')
    
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        export_df.to_excel(writer, index=False, sheet_name='Results')
        workbook = writer.book
        worksheet = writer.sheets['Results']
        
        link_fmt = workbook.add_format({'font_color': 'blue', 'underline': 1})
        
        pmid_idx = export_df.columns.get_loc("PMID") if "PMID" in export_df.columns else -1
        pmcid_idx = export_df.columns.get_loc("PMCID") if "PMCID" in export_df.columns else -1
        doi_idx = export_df.columns.get_loc("DOI") if "DOI" in export_df.columns else -1
        
        # We ensure they exist before trying to zip, but safe zip handles it if length is same
        pmid_col = export_df['PMID'] if "PMID" in export_df.columns else [''] * len(export_df)
        pmcid_col = export_df['PMCID'] if "PMCID" in export_df.columns else [''] * len(export_df)
        doi_col = export_df['DOI'] if "DOI" in export_df.columns else [''] * len(export_df)
        
        for row_num, (pmid_url, pmcid_url, doi_url) in enumerate(zip(pmid_col, pmcid_col, doi_col), start=1):
            if pmid_url and pmid_idx != -1:
                try:
                    display_id = pmid_url.strip("/").split("/")[-1]
                except:
                    display_id = "Link"
                worksheet.write_url(row_num, pmid_idx, pmid_url, string=display_id, cell_format=link_fmt)
                
            if pmcid_url and pmcid_idx != -1:
                try:
                    display_id = pmcid_url.strip("/").split("/")[-1]
                except:
                    display_id = "Link"
                worksheet.write_url(row_num, pmcid_idx, pmcid_url, string=display_id, cell_format=link_fmt)
            
            if doi_url and doi_idx != -1:
                display_doi = doi_url.replace("https://doi.org/", "")
                worksheet.write_url(row_num, doi_idx, doi_url, string=display_doi, cell_format=link_fmt)
                
        worksheet.autofit()
        
    return output.getvalue()

def generate_word_summary(pmid_urls):
    """Fetches abstracts for selected PMIDs (passed as URLs) and creates a Word doc."""
    Entrez.email = "pubmed_tool_web@example.com"
    doc = Document()
    doc.add_heading('PubMed Article Summaries', level=0)
    
    # Extract IDs from URLs
    clean_ids = []
    for url in pmid_urls:
        if url and "pubmed" in url:
            parts = url.strip("/").split("/")
            if parts:
                clean_ids.append(parts[-1])
    
    if not clean_ids:
        return None

    try:
        handle = Entrez.efetch(db="pubmed", id=",".join(clean_ids), rettype="abstract", retmode="xml")
        records = Entrez.read(handle)
        handle.close()
        
        for record in records.get("PubmedArticle", []):
            citation = record.get("MedlineCitation", {})
            article = citation.get("Article", {})
            pmid = citation.get("PMID", "N/A")
            title = article.get("ArticleTitle", "No title")
            
            authors_list = []
            for author in article.get("AuthorList", []):
                if "LastName" in author and "ForeName" in author:
                    authors_list.append(f"{author['ForeName']} {author['LastName']}")
            authors_str = ", ".join(authors_list) if authors_list else "No authors listed"

            journal = article.get("Journal", {}).get("Title", "No journal")
            
            abstract_parts = article.get("Abstract", {}).get("AbstractText", [])
            abstract_text = " ".join(abstract_parts) if abstract_parts else "No abstract found."

            doc.add_heading(f"PMID: {pmid}", level=2)
            
            p = doc.add_paragraph()
            p.add_run("Title: ").bold = True
            p.add_run(title)
            
            p = doc.add_paragraph()
            p.add_run("Authors: ").bold = True
            p.add_run(authors_str)
            
            p = doc.add_paragraph()
            p.add_run("Journal: ").bold = True
            p.add_run(journal)
            
            p = doc.add_paragraph()
            p.add_run("Abstract: ").bold = True
            p.add_run(abstract_text)
            
            doc.add_paragraph("_" * 50) 

        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        return doc_buffer.getvalue()

    except Exception as e:
        print(e)
        return None

col_header, col_tutorial = st.columns([7, 1]) # Adjust ratio to move button

with col_header:
    st.title("🧬 PubMed Research Scanner")
    st.markdown("Search PubMed, select articles, and download Excel lists or Word summaries.")

with col_tutorial:
    st.link_button("Turkish Video Tutorial", "https://www.youtube.com/watch?v=KvsBj1QGqso")

with st.sidebar:
    st.header("Configuration")
    uploaded_scimago = st.file_uploader("Upload Scimago CSV (Optional)", type=["csv"])
    
    scimago_source = None
    default_filename = "scimago.csv"
    
    if uploaded_scimago is not None:
        scimago_source = uploaded_scimago
        st.success("✅ Using your uploaded CSV.")
    elif os.path.exists(default_filename):
        scimago_source = default_filename
        st.info("ℹ️ Using default 'scimago.csv'.")
    else:
        st.warning("⚠️ No Scimago file found.")

col1, col2 = st.columns(2)
with col1:
    kw_or = st.text_input("OR Keywords (e.g. lung cancer, nsclc)")
    kw_and = st.text_input("AND Keywords (e.g. biomarker)")
    study_type = st.multiselect("Study Types", STUDY_TYPES)
with col2:
    start_year = st.text_input("Start Year", value="2020")
    end_year = st.text_input("End Year", value="2025")
    max_results = st.number_input("Max Results", min_value=10, max_value=5000, value=50)

k_or_list = [x.strip() for x in kw_or.split(",") if x.strip()]
k_and_list = [x.strip() for x in kw_and.split(",") if x.strip()]
or_part = " OR ".join([f'"{kw}"[Title/Abstract]' for kw in k_or_list])
and_part = " AND ".join([f'"{kw}"[Title/Abstract]' for kw in k_and_list])
type_part = " OR ".join([f'"{t}"[Publication Type]' for t in study_type])
date_part = f'("{start_year}"[Date - Publication] : "{end_year}"[Date - Publication])'

parts = []
if or_part: parts.append(f"({or_part})")
if and_part: parts.append(f"({and_part})")
if type_part: parts.append(f"({type_part})")
parts.append(date_part)
final_query = " AND ".join(parts)

if 'search_results' not in st.session_state:
    st.session_state.search_results = pd.DataFrame()

if st.button("🔎 Start Search", type="primary"):
    with st.spinner("Searching PubMed..."):
        df = search_pubmed(final_query, max_results)
        
        if df.empty:
            st.warning("No results found.")
            st.session_state.search_results = pd.DataFrame()
        else:
            if scimago_source:
                df = process_quartiles(df, scimago_source)
            else:
                df["Quartile"] = "Unknown (No File)"
            
            cols = ["Select", "PMID", "PMCID", "Corresp. Author Name", "Corresp. Author Email", "Quartile", "Title", "First Author", "Journal", "Year", "DOI", "Article Type", "Keywords"]
            df = df[cols]
            st.session_state.search_results = df

if not st.session_state.search_results.empty:
    st.divider()
    st.subheader("Search Results")
    st.caption("Select rows to generate a Word summary.")

    edited_df = st.data_editor(
        st.session_state.search_results,
        column_config={
            "Select": st.column_config.CheckboxColumn(
                "Select",
                help="Select to include in Word Summary",
                default=False,
            ),
            "PMID": st.column_config.LinkColumn(
                label="PMID",
                display_text=r"https://pubmed\.ncbi\.nlm\.nih\.gov/(.*?)/"
            ),
            "PMCID": st.column_config.LinkColumn(
                label="PMCID",
                display_text=r"https://www\.ncbi\.nlm\.nih\.gov/pmc/articles/(.*?)/"
            ),
            "DOI": st.column_config.LinkColumn(
                label="DOI",
                display_text=r"https://doi\.org/(.*)"
            )
        },
        disabled=["PMID", "PMCID", "Corresp. Author Name", "Corresp. Author Email", "Quartile", "Title", "First Author", "Journal", "Year", "DOI", "Article Type", "Keywords"],
        hide_index=True,
        use_container_width=True
    )

    col_d1, col_d2 = st.columns([1, 1])

    with col_d1:
        excel_data = to_excel(edited_df)
        st.download_button(
            label="📥 Download Excel List",
            data=excel_data,
            file_name="PubMed_List.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )

    with col_d2:
        selected_rows = edited_df[edited_df["Select"] == True]
        
        if not selected_rows.empty:
            if st.button("📄 Generate Word Summary for Selected"):
                with st.spinner("Fetching abstracts and generating Word doc..."):
                    pmid_urls = selected_rows["PMID"].astype(str).tolist()
                    word_data = generate_word_summary(pmid_urls)
                    
                    if word_data:
                        st.download_button(
                            label="⬇️ Download Word Summary (.docx)",
                            data=word_data,
                            file_name="PubMed_Abstracts_Summary.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                    else:
                        st.error("Failed to generate Word document.")
        else:
            st.info("Select checkboxes above to enable Word summary generation.")
